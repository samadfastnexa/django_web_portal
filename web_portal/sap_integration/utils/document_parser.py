"""
Utility to parse Word documents dynamically with different format handling
Parse product description documents from media/product_images folders
"""
import os
import mammoth
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
import base64
from bs4 import BeautifulSoup
import logging

logger = logging.getLogger(__name__)


class WordDocumentParser:
    """Parse Word documents and preserve formatting dynamically"""
    
    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = None
        self.images = []
        
        if os.path.exists(file_path):
            try:
                self.doc = Document(file_path)
            except Exception as e:
                logger.error(f"Error loading document {file_path}: {e}")
    
    def parse_to_html_mammoth(self):
        """Convert Word document to HTML using mammoth (best for complex formatting)"""
        try:
            with open(self.file_path, 'rb') as doc_file:
                result = mammoth.convert_to_html(
                    doc_file,
                    style_map="""
                        p[style-name='Heading 1'] => h1.heading1:fresh
                        p[style-name='Heading 2'] => h2.heading2:fresh
                        p[style-name='Heading 3'] => h3.heading3:fresh
                        p[style-name='Heading 4'] => h4.heading4:fresh
                        p[style-name='Title'] => h1.title:fresh
                        table => table.product-table
                        r[style-name='Strong'] => strong
                        r[style-name='Emphasis'] => em
                    """,
                    convert_image=mammoth.images.img_element(self._image_converter)
                )
                
                # Add RTL wrapper
                html = f'<div class="product-document rtl-content">{result.value}</div>'
                return html, result.messages
                
        except Exception as e:
            logger.error(f"Error parsing document with mammoth: {e}")
            return None, [str(e)]
    
    def _image_converter(self, image):
        """Convert images to base64 data URIs"""
        try:
            with image.open() as image_bytes:
                img_data = image_bytes.read()
                encoded = base64.b64encode(img_data).decode('utf-8')
                content_type = image.content_type or 'image/png'
                return {"src": f"data:{content_type};base64,{encoded}"}
        except Exception as e:
            logger.error(f"Error converting image: {e}")
            return {"src": ""}
    
    def parse_custom_formatting(self):
        """Custom parser for complex formatting preservation"""
        if not self.doc:
            return None
        
        html_parts = ['<div class="product-document rtl-content">']
        
        # Process paragraphs
        for paragraph in self.doc.paragraphs:
            html_parts.append(self._parse_paragraph(paragraph))
        
        # Process tables
        for table in self.doc.tables:
            html_parts.append(self._parse_table(table))
        
        html_parts.append('</div>')
        
        return '\n'.join(html_parts)
    
    def _parse_paragraph(self, paragraph):
        """Parse paragraph with all formatting"""
        text = paragraph.text.strip()
        
        if not text:
            return ''
        
        # Determine tag based on style
        style_name = paragraph.style.name.lower()
        
        if 'heading 1' in style_name or 'title' in style_name:
            tag = 'h1'
        elif 'heading 2' in style_name:
            tag = 'h2'
        elif 'heading 3' in style_name:
            tag = 'h3'
        elif 'heading 4' in style_name:
            tag = 'h4'
        else:
            tag = 'p'
        
        # Get alignment
        alignment = self._get_alignment(paragraph)
        
        # Build inline styles
        styles = []
        
        # Check for custom colors, fonts in runs
        runs_html = []
        for run in paragraph.runs:
            run_style = self._get_run_style(run)
            run_text = run.text
            
            if run_style:
                runs_html.append(f'<span style="{run_style}">{run_text}</span>')
            else:
                runs_html.append(run_text)
        
        if alignment:
            styles.append(f'text-align: {alignment}')
        
        style_attr = f' style="{"; ".join(styles)}"' if styles else ''
        
        return f'<{tag}{style_attr}>{"".join(runs_html)}</{tag}>'
    
    def _get_run_style(self, run):
        """Extract style from a run (text fragment)"""
        styles = []
        
        # Font color
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            color = f'#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}'
            styles.append(f'color: {color}')
        
        # Font size
        if run.font.size:
            size_pt = run.font.size.pt
            styles.append(f'font-size: {size_pt}pt')
        
        # Font name
        if run.font.name:
            styles.append(f'font-family: "{run.font.name}"')
        
        # Bold
        if run.bold:
            styles.append('font-weight: bold')
        
        # Italic
        if run.italic:
            styles.append('font-style: italic')
        
        # Underline
        if run.underline:
            styles.append('text-decoration: underline')
        
        # Background color
        if run.font.highlight_color:
            styles.append(f'background-color: {run.font.highlight_color}')
        
        return '; '.join(styles) if styles else None
    
    def _get_alignment(self, paragraph):
        """Get paragraph alignment"""
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: 'left',
            WD_ALIGN_PARAGRAPH.CENTER: 'center',
            WD_ALIGN_PARAGRAPH.RIGHT: 'right',
            WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
        }
        
        if paragraph.alignment:
            return alignment_map.get(paragraph.alignment, None)
        return None
    
    def _parse_table(self, table):
        """Convert Word table to HTML with formatting"""
        html = ['<table class="product-table table table-bordered">']
        
        for i, row in enumerate(table.rows):
            html.append('<tr>')
            
            for j, cell in enumerate(row.cells):
                # Check if it's a header row (first row or bold text)
                is_header = i == 0
                tag = 'th' if is_header else 'td'
                
                # Get cell background color
                cell_style = self._get_cell_style(cell)
                style_attr = f' style="{cell_style}"' if cell_style else ''
                
                # Get cell content with formatting
                cell_html = []
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run_style = self._get_run_style(run)
                        if run_style:
                            cell_html.append(f'<span style="{run_style}">{run.text}</span>')
                        else:
                            cell_html.append(run.text)
                
                content = ''.join(cell_html) or cell.text
                html.append(f'<{tag}{style_attr}>{content}</{tag}>')
            
            html.append('</tr>')
        
        html.append('</table>')
        return '\n'.join(html)
    
    def _get_cell_style(self, cell):
        """Get table cell background color"""
        styles = []
        
        try:
            # Get cell shading/background
            if hasattr(cell, '_element') and hasattr(cell._element, 'tcPr'):
                tc_pr = cell._element.tcPr
                if tc_pr is not None:
                    shading = tc_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                    if shading is not None and shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill'):
                        color = shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                        if color and color != 'auto':
                            styles.append(f'background-color: #{color}')
        except Exception as e:
            pass
        
        return '; '.join(styles) if styles else None
    
    def get_document_info(self):
        """Get document metadata"""
        if not self.doc:
            return {}
        
        return {
            'paragraphs': len(self.doc.paragraphs),
            'tables': len(self.doc.tables),
            'sections': len(self.doc.sections),
        }


def parse_product_document(file_path, method='mammoth'):
    """
    Parse product document with specified method
    
    Args:
        file_path: Path to the Word document
        method: 'mammoth' for automatic conversion or 'custom' for detailed parsing
    
    Returns:
        HTML string with formatted content
    """
    if not os.path.exists(file_path):
        return f'<div class="alert alert-warning rtl-content">دستاویز فائل نہیں ملی: {file_path}</div>'
    
    parser = WordDocumentParser(file_path)
    
    if method == 'mammoth':
        html, messages = parser.parse_to_html_mammoth()
        if html:
            return html
        else:
            # Fallback to custom parser
            return parser.parse_custom_formatting()
    else:
        return parser.parse_custom_formatting()


def get_product_document_path(product_urdu_name, product_urdu_ext, database_name=''):
    """
    Get the full path to a product description document
    
    Args:
        product_urdu_name: Document filename (without extension)
        product_urdu_ext: Document file extension
        database_name: Database/schema name from Company model (e.g., 4B-BIO_APP, 4B-ORANG_APP, 4B-AGRI_LIVE)
    
    Returns:
        Full file path or None if not found
    """
    from django.conf import settings
    import glob
    
    media_root = settings.MEDIA_ROOT
    product_images_dir = os.path.join(media_root, 'product_images')
    
    # Dynamically determine possible folders to search
    possible_folders = []
    
    if database_name:
        # Extract folder name from database name by removing common suffixes
        # Examples: 4B-BIO_APP -> 4B-BIO, 4B-ORANG_APP -> 4B-ORANG, 4B-AGRI_LIVE -> 4B-AGRI
        folder_name = database_name.replace('_APP', '').replace('_LIVE', '').replace('_TEST', '').strip()
        
        # If extracted folder exists, prioritize it
        folder_path = os.path.join(product_images_dir, folder_name)
        if os.path.exists(folder_path) and os.path.isdir(folder_path):
            possible_folders.append(folder_name)
    
    # Get all available folders in product_images directory dynamically
    try:
        all_folders = [d for d in os.listdir(product_images_dir) 
                      if os.path.isdir(os.path.join(product_images_dir, d)) and not d.startswith('.')]
        # Add any folders not already in the list
        for folder in all_folders:
            if folder not in possible_folders:
                possible_folders.append(folder)
    except Exception:
        # Fallback: if we can't read directory, try common patterns
        if not possible_folders:
            possible_folders = ['4B-BIO', '4B-ORANG', '4B-AGRI']
    
    # Try each folder
    for folder in possible_folders:
        # Try with provided extension
        file_path = os.path.join(product_images_dir, folder, f'{product_urdu_name}.{product_urdu_ext}')
        
        if os.path.exists(file_path):
            return file_path
        
        # Try alternative extensions
        base_path = os.path.join(product_images_dir, folder, product_urdu_name)
        for ext in ['docx', 'doc', 'DOCX', 'DOC', 'pdf', 'PDF']:
            alt_path = f'{base_path}.{ext}'
            if os.path.exists(alt_path):
                return alt_path
    
    return None
